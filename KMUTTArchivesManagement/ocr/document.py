import os
import docx
import re
from io import BytesIO
from PIL import Image

### Global variable ###

ROOT = os.path.abspath(os.getcwd())
PATH_REPORT = os.path.join(ROOT, 'document-report')
REGEX = re.compile(r"[^\u0E00-\u0E7Fa-zA-Z'\d\\/\.\-\s]")

########################


def checkFile(pathIn):
    # check file and directory it exitst?
    if os.path.exists(pathIn):
        return True
    return False


def createDirectory(path):
    # check path and then create directory
    if not checkFile(path):
        try:
            os.mkdir(path)
        except OSError:
            print("Creation of the directory %s failed" % path)
    # else:
        # print("Already have directory %s" % path)


def createDoc(path):
    # check file and then create doc or open it if already created
    if not checkFile(path):
        try:
            mydoc = docx.Document()
            print("Successfully created the file %s" % path)
        except OSError:
            print("Creation of the directory %s failed" % path)
    else:
        mydoc = docx.Document(path)
        print("Successfully open the file %s" % path)
    # mydoc.save(path)
    return mydoc


def addReportDocText(text, doc, clean=False):
    if not clean:
        text = "Full text:"+text
    doc.add_paragraph(text)


def addReportDocPicture(picture, doc):
    # save temp file picture
    image = Image.fromarray(picture)
    bytesImage = BytesIO()
    image.save(bytesImage, format="PNG")
    # insert picture in to doc
    doc.add_picture(bytesImage, width=docx.shared.Inches(2))


def saveTempImg(picture):
    image = Image.fromarray(picture)
    bytesImage = BytesIO()
    image.save(bytesImage, format="PNG")
    return bytesImage


def addReportDoc(arrayPicture, arrayText, arrayCleanText, path):
    if len(arrayText) != len(arrayPicture):
        print('length between arrayText and arrayPicture is not equal!!!!!!!')
        return
    doc = createDoc(path)
    for inx in range(len(arrayText)):
        addReportDocPicture(arrayPicture[inx], doc)
        addReportDocText(arrayText[inx], doc)
        doc.add_paragraph('Clean text:')
        # for text in arrayCleanText[inx]:
        addReportDocText(arrayCleanText[inx], doc, True)
    doc.save(path)


def addReportText(fullText, page, fileName):
    # create file text prepare sent to next step
    with open(PATH_REPORT+"/"+fileName+"/"+fileName+"page-"+str(page)+".txt", 'w', encoding='utf8') as f:
        f.write(fullText)


def cleanTextRegex(fullText):
    # remove anything except TH/EN alphabet, digit, '/'
    word = b'\x0c'
    word = word.decode('utf8')
    text = fullText.replace(word, "")
    newText = ''
    for index, current in enumerate(text):
        if not(current == ' ' and re.match('[A-Za-z]', text[index - 1]) == None and re.match('[A-Za-z]', text[index + 1]) == None):
            newText += current
    newText = REGEX.sub('', newText)
    if text:
        return text + "\n"
    return False
